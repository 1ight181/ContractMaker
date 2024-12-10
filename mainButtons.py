from docx import Document
from docxtpl import DocxTemplate
from PyQt5.QtWidgets import QListWidget, QPushButton, QComboBox, QFileDialog, QAction, QLineEdit, QTextEdit
from settings import SettingsWindow
from config import Config
from callMessageBox import *
import re, datetime


def openSettingsWindow(parent: QtWidgets.QMainWindow):
    settingsWindow = SettingsWindow(parent)

def getNameOfOrder():
    dlg = QFileDialog()
    dlg.setFileMode(QFileDialog.FileMode.AnyFile)
    dlg.setNameFilter("Word file (*.docx *.doc)")
    dlg.setAcceptMode(QFileDialog.AcceptMode.AcceptOpen)

    try:
        if dlg.exec():
            fileName = dlg.selectedFiles()
            return fileName[0]
    except:
        return None
    

def getPathToDirContracts():
    dlg = QFileDialog()
    return dlg.getExistingDirectory(caption="Выберите папку, где будут созданы договора")
    

def openOrderButton(listNumberOfStudent: QListWidget,
                    listNameOfStudent: QListWidget, 
                    listGroupOfStudent: QListWidget, 
                    listPlaceOfPractice: QListWidget, 
                    listHeadOfPractice: QListWidget, 
                    comboBoxDirectionOfStudy:QComboBox, 
                    comboBoxTypeOfPractice:QComboBox, 
                    comboBoxDurationOfPractice:QComboBox,
                    lineEditNumberOfAuthority: QLineEdit,
                    lineEditDateOfAuthority: QLineEdit,
                    pushButtonCreateContract: QPushButton, 
                    actionCreateContract:QAction,
                    textEditPeriodOfPractice: QTextEdit,
                    nameOfOrganization: str,
                    isCreateContractForOrganization: bool,
                    directionsOfStudyParagraph: int,
                    typesOfPracticeParagraph: int,
                    periodOfPracticeParagraph: int
                    ):
    
    findSpacesRegEx = re.compile('\s+')

    fileName = getNameOfOrder()
    if not fileName:
        return 

    listNameOfStudent.clear()
    listGroupOfStudent.clear()
    listPlaceOfPractice.clear()
    listHeadOfPractice.clear()
    listHeadOfPractice.clear()
    listNumberOfStudent.clear()
    
    textEditPeriodOfPractice.clear()

    comboBoxTypeOfPractice.setCurrentText("Не выбранно")
    comboBoxDurationOfPractice.setCurrentText("Не выбранно")
    comboBoxDirectionOfStudy.setCurrentText("Не выбранно")

    document = Document(fileName)

    tableOfstudents = document.tables[0]
    textOfRowDict = dict.fromkeys(["№",
                                  "Фамилия имя отчество (при наличии) полностью)",
                                  "Академическая группа",
                                  "Наименование места прохождения практической подготовки (организационно-правовая форма в аббревиатуре)",
                                  "Руководитель практики (должность, фамилия имя отчество (при наличии) полностью)"])

    for i, row in enumerate(tableOfstudents.rows):
        j = i+1
        for j, cell in enumerate(row.cells):
            if i == 0:
                if re.sub(findSpacesRegEx, " ", cell.text) == "Наименование места прохождения практической подготовки (организационно-правовая форма в аббревиатуре)":
                    numberOfCellOfUust = j 
                continue
            row.cells[numberOfCellOfUust].text

            if isCreateContractForOrganization:
                if row.cells[numberOfCellOfUust].text == nameOfOrganization:        
                    continue

            textOfRowDict[re.sub(findSpacesRegEx, " ", tableOfstudents.rows[0].cells[j].text)] = re.sub(findSpacesRegEx, " ", cell.text)

        listNameOfStudent.addItem(textOfRowDict.get("Фамилия имя отчество (при наличии) полностью)"))
        listGroupOfStudent.addItem(textOfRowDict.get("Академическая группа"))
        listPlaceOfPractice.addItem(textOfRowDict.get("Наименование места прохождения практической подготовки (организационно-правовая форма в аббревиатуре)"))
        listHeadOfPractice.addItem(textOfRowDict.get("Руководитель практики (должность, фамилия имя отчество (при наличии) полностью)"))
        textOfRowDict.clear()
    
    for i in range(listNameOfStudent.count()):
        listNumberOfStudent.addItem(str(i+1))
    i = 0

    directionsOfStudy = [comboBoxDirectionOfStudy.itemText(i) for i in range(comboBoxDirectionOfStudy.count())]

    for direction in directionsOfStudy:
        if re.search(re.sub(findSpacesRegEx, " ", direction), document.paragraphs[directionsOfStudyParagraph].text + r" "):
            comboBoxDirectionOfStudy.setCurrentText(direction)

    typesOfPractice = [comboBoxTypeOfPractice.itemText(i) for i in range(comboBoxTypeOfPractice.count())]

    for typeOfPractice in typesOfPractice:
        listOfMatches = [""]
        listOfMatches += re.findall(typeOfPractice[0:4].lower(), document.paragraphs[typesOfPracticeParagraph].text)

        if listOfMatches[-1] == typeOfPractice[0:4].lower():
            comboBoxTypeOfPractice.setCurrentText(typeOfPractice)
            comboBoxDurationOfPractice.setCurrentIndex(comboBoxTypeOfPractice.currentIndex())
    
    listOfMatches = []
    listOfMatches += re.findall(r"\d{2}\.\d{2}\.\d{4}", document.paragraphs[periodOfPracticeParagraph].text)

    resultString = ""

    i = 0
    while(i < len(listOfMatches)-1):
        resultString += listOfMatches[i] + " - " + listOfMatches[i+1] + "<br />"
        i += 2

    textEditPeriodOfPractice.setHtml(resultString)

    pushButtonCreateContract.setEnabled(True)
    comboBoxDirectionOfStudy.setEnabled(True)
    comboBoxTypeOfPractice.setEnabled(True)
    comboBoxDurationOfPractice.setEnabled(True)
    actionCreateContract.setEnabled(True)
    lineEditNumberOfAuthority.setEnabled(True)
    lineEditDateOfAuthority.setEnabled(True)
    textEditPeriodOfPractice.setEnabled(True)

def iterAllItems(qlist: QListWidget):
    for i in range(qlist.count()):
        yield qlist.item(i).text()

def saveContractButton (
                        listNameOfStudent: QListWidget,
                        listGroupOfStudent: QListWidget, 
                        listPlaceOfPractice: QListWidget, 
                        comboBoxDirectionOfStudy: QComboBox, 
                        comboBoxTypeOfPractice: QComboBox, 
                        comboBoxDurationOfPractice: QComboBox,
                        lineEditNumberOfAuthority: QLineEdit,
                        lineEditDateOfAuthority: QLineEdit,
                        textEditPeriodOfPractice: QTextEdit
                        ):
    
    if comboBoxDirectionOfStudy.currentText() == "" or comboBoxTypeOfPractice.currentText() == "" or comboBoxDurationOfPractice.currentText() == "":
        callErrorMessageBox("Откройте приказ снова")
        return
    
    if lineEditNumberOfAuthority.text() == "" or lineEditDateOfAuthority.text() == "":
        callErrorMessageBox("Заполните поле номера и даты корректными значениями")
        return

    try:
        datetime.datetime.strptime(lineEditDateOfAuthority.text(), "%d.%m.%Y")
    except ValueError:
        callErrorMessageBox("Заполните поле даты корректным значением в формате ДД.ММ.ГГГГ")
        return
    
    if re.sub(r"\s*((0?[1-9]|[12][0-9]|3[01])\.(0?[1-9]|1[012])\.\d{4} *-+ *(0?[1-9]|[12][0-9]|3[01])\.(0?[1-9]|1[012])\.\d{4})\,?\s*", "", textEditPeriodOfPractice.toPlainText()) != "":
        callErrorMessageBox("Заполните поле срока практики корректными значениями в формате ДД.ММ.ГГГГ - ДД.ММ.ГГГГ, " 
                            + "если периодов несколько отделите их с помощью клавиши ENTER, либо запятой")
        return

    newDirName = getPathToDirContracts()

    if not newDirName:
        return

    indexes = {}
    contextNameOfStudents = []

    for index in range(listPlaceOfPractice.count()):
        indexes.setdefault(index+1, listPlaceOfPractice.item(index).text())
    
    tuplePlacesOfPractice = tuple({item for item in iterAllItems(listPlaceOfPractice)})
    i = 0
    k = 0
    for place in tuplePlacesOfPractice:
        i += 1
        newDocument = DocxTemplate("contract example.docx")

        for j in [index for index in indexes if indexes[index] == place]:
            k += 1
            print(j-1)
            contextNameOfStudents.append(str(k)+ ". " + listNameOfStudent.item(j-1).text())
        try:
            newDocument.render({"studentName": contextNameOfStudents,
                            "placeOfPractice": place,
                            "number": 1,
                            "directionOfStudy": comboBoxDirectionOfStudy.currentText(),
                            "typeOfPractice": comboBoxTypeOfPractice.currentText(),
                            "course": re.findall(r'\d', listGroupOfStudent.item(1).text())[0],
                            "countOfStudents": len(contextNameOfStudents),
                            "durationOfPractice": comboBoxDurationOfPractice.currentText(),
                            "numberOfAuthority": lineEditNumberOfAuthority.text(),
                            "dateOfAuthority": lineEditDateOfAuthority.text(),
                            "periodOfPractice": "\n".join(re.findall(r"\d{2}\.\d{2}\.\d{4} *-+ *\d{2}\.\d{2}\.\d{4}", textEditPeriodOfPractice.toPlainText()))})
            newDocument.save(newDirName.replace("\\", "\\\\") + "/" + place + ".docx")
        except PermissionError:
            callErrorMessageBox("Закройте открытый договор!")
            return
            
        contextNameOfStudents = []
        k = 0 
    
    callInfoMessageBox("Договора находятся в папке, которую вы указали")

